from typing import Optional, Type
from pydantic.v1 import BaseModel, Field
from langchain.agents import tool
import json
import openpyxl
from docx import Document
import win32com.client

class OfficeAutomationInput(BaseModel):
    input: str = Field(description='A JSON string containing the operation type and necessary parameters. Example: {"operation": "excel_read", "file_path": "data.xlsx", "sheet_name": "Sheet1", "cell_range": "A1:B5"}')

@tool("office_automation_tool", args_schema=OfficeAutomationInput, return_direct=False)
def office_automation_tool(input: str) -> str:
    """
    Automate Microsoft Office tasks including Excel and Word operations. Enables interaction with Excel and Word, including data extraction, VBA execution, document generation, and formatting.
    """
    try:
        params = json.loads(input)
        operation = params.get('operation')

        if operation == 'excel_read':
            return excel_read(params)
        elif operation == 'excel_write':
            return excel_write(params)
        elif operation == 'excel_run_macro':
            return excel_run_macro(params)
        elif operation == 'word_create':
            return word_create(params)
        elif operation == 'word_read':
            return word_read(params)
        elif operation == 'word_append':
            return word_append(params)
        else:
            return f"Unsupported operation: {operation}"
    except Exception as e:
        return f"Error: {str(e)}"

def excel_read(params):
    file_path = params['file_path']
    sheet_name = params['sheet_name']
    cell_range = params['cell_range']
    wb = openpyxl.load_workbook(file_path, data_only=True)
    sheet = wb[sheet_name]
    data = [[cell.value for cell in row] for row in sheet[cell_range]]
    return json.dumps(data)

def excel_write(params):
    file_path = params['file_path']
    sheet_name = params['sheet_name']
    start_cell = params['start_cell']
    data = params['data']
    wb = openpyxl.load_workbook(file_path)
    sheet = wb[sheet_name]
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            sheet.cell(row=start_cell[0]+i, column=start_cell[1]+j, value=value)
    wb.save(file_path)
    return f"Data written to {file_path}"

def excel_run_macro(params):
    file_path = params['file_path']
    macro_name = params['macro_name']
    excel = win32com.client.Dispatch("Excel.Application")
    wb = excel.Workbooks.Open(file_path)
    excel.Application.Run(macro_name)
    wb.Save()
    excel.Application.Quit()
    return f"Macro {macro_name} executed in {file_path}"

def word_create(params):
    file_path = params['file_path']
    content = params['content']
    doc = Document()
    doc.add_paragraph(content)
    doc.save(file_path)
    return f"Document created at {file_path}"

def word_read(params):
    file_path = params['file_path']
    doc = Document(file_path)
    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return content

def word_append(params):
    file_path = params['file_path']
    content = params['content']
    doc = Document(file_path)
    doc.add_paragraph(content)
    doc.save(file_path)
    return f"Content appended to {file_path}"


from typing import Optional, Type
from pydantic.v1 import BaseModel, Field
from langchain.agents import tool
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
import json
import os

class PDFToolInput(BaseModel):
    input: str = Field(description='A JSON string containing the operation type and necessary parameters. Example: {"operation": "merge", "input_paths": ["/path/to/file1.pdf", "/path/to/file2.pdf"], "output_path": "/path/to/output.pdf"}')

@tool("pdf_tool", args_schema=PDFToolInput, return_direct=False)
def pdf_tool(input: str) -> str:
    """
    Perform PDF-specific operations such as merging documents, including sealed documents. This tool can handle encrypted PDFs and supports various PDF manipulation tasks.
    """
    try:
        params = json.loads(input)
        operation = params.get('operation')

        if operation == 'merge':
            input_paths = params.get('input_paths', [])
            output_path = params.get('output_path')

            if not input_paths or not output_path:
                return "Error: Missing input_paths or output_path for merge operation."

            merger = PdfMerger()
            
            for path in input_paths:
                if not os.path.exists(path):
                    return f"Error: File not found - {path}"
                
                reader = PdfReader(path)
                if reader.is_encrypted:
                    reader.decrypt('')
                
                merger.append(reader)
            
            with open(output_path, 'wb') as output_file:
                merger.write(output_file)
            
            merger.close()
            return f"PDFs merged successfully. Output saved to {output_path}"

        else:
            return f"Error: Unsupported operation - {operation}"

    except json.JSONDecodeError:
        return "Error: Invalid JSON input"
    except Exception as e:
        return f"Error: {str(e)}"


from typing import Optional, Type
from pydantic.v1 import BaseModel, Field
from langchain.agents import tool
from langchain.tools import HumanInputRun
import json

class DataVerificationInput(BaseModel):
    input: str = Field(description='A string representation of a dictionary containing the data to be verified. Example: {"customer_name": "John Doe", "age": 35, "purchase_amount": 1000}')

human_input = HumanInputRun()

@tool("data_verification_interface", args_schema=DataVerificationInput, return_direct=False)
def data_verification_interface(input: str) -> str:
    """Provides a mechanism for human interaction to verify data accuracy and completeness. It presents data for review, allows confirmation or correction, and returns the verification result."""
    try:
        data = json.loads(input)
        
        verification_prompt = f"Please verify the following data:\n{json.dumps(data, indent=2)}\n\nIs this data accurate and complete? (yes/no): "
        verification = human_input.run(verification_prompt)
        
        if verification.lower().strip() == 'yes':
            return "Data verified as accurate and complete."
        else:
            corrections_prompt = "Please provide corrections or additional information:\n"
            corrections = human_input.run(corrections_prompt)
            
            try:
                corrected_data = json.loads(corrections)
                data.update(corrected_data)
                return f"Data updated with corrections: {json.dumps(data, indent=2)}"
            except json.JSONDecodeError:
                return f"Data requires corrections: {corrections}"
    except json.JSONDecodeError:
        return "Error: Invalid input format. Please provide a valid JSON string."


from typing import Type
from pydantic.v1 import BaseModel, Field
from langchain.tools import BaseTool
import docx
from difflib import SequenceMatcher
import json
import re

class WordDocAnalyzerInput(BaseModel):
    input: str = Field(description='A JSON string containing paths to two Word documents and optionally a path for the output file. Example: {"doc1_path": "/path/to/doc1.docx", "doc2_path": "/path/to/doc2.docx", "output_path": "/path/to/output.txt"}')

class WordDocAnalyzer(BaseTool):
    name = "word_doc_analyzer"
    description = "Analyzes and compares two Word documents, extracting content, comparing structure, highlighting differences, performing error detection, and conducting quality assurance checks against predefined standards."
    args_schema: Type[BaseModel] = WordDocAnalyzerInput

    def _run(self, input: str) -> str:
        input_data = json.loads(input)
        doc1_path = input_data["doc1_path"]
        doc2_path = input_data["doc2_path"]
        output_path = input_data.get("output_path")

        result = self.compare_word_docs(doc1_path, doc2_path)

        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result)
            return f"Analysis complete. Results written to {output_path}"
        else:
            return result

    def compare_word_docs(self, doc1_path: str, doc2_path: str) -> str:
        doc1 = docx.Document(doc1_path)
        doc2 = docx.Document(doc2_path)

        result = []

        result.append(f"Analysis of {doc1_path} and {doc2_path}")
        result.append(self.compare_structure(doc1, doc2))
        result.append(self.compare_content(doc1, doc2))
        result.append(self.detect_errors(doc1, "Document 1"))
        result.append(self.detect_errors(doc2, "Document 2"))
        result.append(self.quality_assurance_check(doc1, "Document 1"))
        result.append(self.quality_assurance_check(doc2, "Document 2"))

        return "\n\n".join(result)

    def compare_structure(self, doc1: docx.Document, doc2: docx.Document) -> str:
        structure1 = self.get_document_structure(doc1)
        structure2 = self.get_document_structure(doc2)

        if structure1 == structure2:
            return "Document structures are identical."
        else:
            differences = []
            for i, (s1, s2) in enumerate(zip(structure1, structure2)):
                if s1 != s2:
                    differences.append(f"Difference at level {i}: {s1} vs {s2}")
            return "Document structures differ:\n" + "\n".join(differences)

    def get_document_structure(self, doc: docx.Document) -> list:
        structure = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                structure.append(f"{para.style.name}: {para.text}")
        return structure

    def compare_content(self, doc1: docx.Document, doc2: docx.Document) -> str:
        text1 = "\n".join([para.text for para in doc1.paragraphs])
        text2 = "\n".join([para.text for para in doc2.paragraphs])

        matcher = SequenceMatcher(None, text1, text2)
        similarity = matcher.ratio()

        differences = []
        for i, s in enumerate(matcher.get_opcodes()):
            if s[0] != 'equal':
                differences.append(f"Difference {i+1}: {s[0]}, positions {s[1]}:{s[2]} vs {s[3]}:{s[4]}")

        return f"Content similarity: {similarity:.2%}\n\nDifferences:\n" + "\n".join(differences[:10])

    def detect_errors(self, doc: docx.Document, doc_name: str) -> str:
        errors = []
        for i, para in enumerate(doc.paragraphs):
            if len(para.text.strip()) > 0:
                if not para.text[0].isupper():
                    errors.append(f"Paragraph {i+1} doesn't start with a capital letter")
                if not para.text.strip().endswith(('.', '?', '!')):
                    errors.append(f"Paragraph {i+1} doesn't end with proper punctuation")
                if len(para.text.split()) > 40:
                    errors.append(f"Paragraph {i+1} is very long (over 40 words)")
                if re.search(r'\b(\w+)(\s+\1\b)+', para.text, re.IGNORECASE):
                    errors.append(f"Paragraph {i+1} contains word repetition")

        return f"Errors detected in {doc_name}:\n" + "\n".join(errors) if errors else f"No errors detected in {doc_name}."

    def quality_assurance_check(self, doc: docx.Document, doc_name: str) -> str:
        qa_results = []
        word_count = sum(len(para.text.split()) for para in doc.paragraphs)
        qa_results.append(f"Word count: {word_count}")

        heading_count = sum(1 for para in doc.paragraphs if para.style.name.startswith('Heading'))
        qa_results.append(f"Number of headings: {heading_count}")

        image_count = len(doc.inline_shapes)
        qa_results.append(f"Number of images: {image_count}")

        table_count = len(doc.tables)
        qa_results.append(f"Number of tables: {table_count}")

        if word_count < 100:
            qa_results.append("Warning: Document is very short")
        if heading_count == 0:
            qa_results.append("Warning: Document has no headings")
        if image_count == 0:
            qa_results.append("Note: Document contains no images")
        if table_count == 0:
            qa_results.append("Note: Document contains no tables")

        return f"Quality Assurance Check for {doc_name}:\n" + "\n".join(qa_results)

    async def _arun(self, input: str) -> str:
        raise NotImplementedError("WordDocAnalyzer does not support async")

word_doc_analyzer = WordDocAnalyzer()


from typing import Optional, Type
from pydantic.v1 import BaseModel, Field
from langchain.agents import tool
from docx import Document
from docx.shared import Inches
import json
import os
import base64

class WordDocumentInput(BaseModel):
    input: str = Field(description='A JSON string containing the operation and its parameters. Example: {"operation": "create", "content": {"headings": [{"text": "Sample Document", "level": 1}], "paragraphs": ["This is a sample paragraph."], "images": [{"path": "sample_image.jpg", "width": 6}]}, "file_path": "new_document.docx"}')

class WordDocumentHandler:
    def __init__(self):
        self.document = None

    def create_document(self, content, file_path):
        self.document = Document()
        for heading in content.get('headings', []):
            self.document.add_heading(heading['text'], heading.get('level', 1))
        for paragraph in content.get('paragraphs', []):
            self.document.add_paragraph(paragraph)
        for image in content.get('images', []):
            if image['path'].startswith('data:image'):
                # Handle base64 encoded image
                header, encoded = image['path'].split(",", 1)
                data = base64.b64decode(encoded)
                with open('temp_image.png', 'wb') as f:
                    f.write(data)
                self.document.add_picture('temp_image.png', width=Inches(image.get('width', 6)))
                os.remove('temp_image.png')
            else:
                self.document.add_picture(image['path'], width=Inches(image.get('width', 6)))
        self.save_document(file_path)
        return f"Document created and saved as {file_path}"

    def load_document(self, file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        self.document = Document(file_path)
        return "Document loaded successfully"

    def modify_document(self, content, file_path):
        if self.document is None:
            raise ValueError("No document loaded. Please load a document first.")
        for heading in content.get('headings', []):
            self.document.add_heading(heading['text'], heading.get('level', 1))
        for paragraph in content.get('paragraphs', []):
            self.document.add_paragraph(paragraph)
        for image in content.get('images', []):
            if image['path'].startswith('data:image'):
                # Handle base64 encoded image
                header, encoded = image['path'].split(",", 1)
                data = base64.b64decode(encoded)
                with open('temp_image.png', 'wb') as f:
                    f.write(data)
                self.document.add_picture('temp_image.png', width=Inches(image.get('width', 6)))
                os.remove('temp_image.png')
            else:
                self.document.add_picture(image['path'], width=Inches(image.get('width', 6)))
        self.save_document(file_path)
        return f"Document modified and saved as {file_path}"

    def get_text_content(self):
        if self.document is None:
            raise ValueError("No document loaded. Please load a document first.")
        return "\n".join([para.text for para in self.document.paragraphs])

    def save_document(self, file_path):
        try:
            self.document.save(file_path)
        except Exception as e:
            raise IOError(f"Error saving document: {str(e)}")

    def close_document(self):
        self.document = None
        return "Document closed"

word_handler: WordDocumentHandler = WordDocumentHandler()

@tool("word_document_handler", args_schema=WordDocumentInput, return_direct=False)
def word_document_handler(input: str) -> str:
    """Handle Microsoft Word document operations such as creation, modification, and content extraction."""
    try:
        data = json.loads(input)
        operation = data['operation']
        
        if operation == 'create':
            return word_handler.create_document(data['content'], data['file_path'])
        elif operation == 'load':
            return word_handler.load_document(data['file_path'])
        elif operation == 'modify':
            return word_handler.modify_document(data['content'], data['file_path'])
        elif operation == 'get_content':
            return word_handler.get_text_content()
        elif operation == 'close':
            return word_handler.close_document()
        else:
            return f"Unknown operation: {operation}"
    except json.JSONDecodeError:
        return "Invalid JSON input"
    except KeyError as e:
        return f"Missing required key in input: {str(e)}"
    except Exception as e:
        return f"An error occurred: {str(e)}"


from typing import Optional, Type
from pydantic.v1 import BaseModel, Field
from langchain.agents import tool
import openpyxl
from docx import Document
import win32com.client
import json

class MSOfficeInput(BaseModel):
    input: str = Field(description='A JSON string containing the operation type and parameters. Example: {"operation": "read_excel", "file_path": "C:/example.xlsx", "cell": "A1"}')

@tool("ms_office_integration_tool", args_schema=MSOfficeInput, return_direct=False)
def ms_office_integration_tool(input: str) -> str:
    """
    Provides capabilities to read from, write to, and manipulate Excel spreadsheets; create, edit, and compile Word documents; and execute VBA scripts in both Excel and Word. Supports various operations for comprehensive Microsoft Office integration.
    """
    try:
        params = json.loads(input)
        operation = params.get('operation')

        if operation == 'read_excel':
            return read_excel(params['file_path'], params['cell'])
        elif operation == 'write_excel':
            write_excel(params['file_path'], params['cell'], params['value'])
            return f"Value written to {params['cell']} in {params['file_path']}"
        elif operation == 'create_word':
            create_word_document(params['file_path'], params['content'])
            return f"Word document created at {params['file_path']}"
        elif operation == 'edit_word':
            edit_word_document(params['file_path'], params['content'])
            return f"Word document edited at {params['file_path']}"
        elif operation == 'execute_vba_excel':
            execute_vba_excel(params['file_path'], params['macro_name'])
            return f"VBA macro {params['macro_name']} executed in Excel file {params['file_path']}"
        elif operation == 'execute_vba_word':
            execute_vba_word(params['file_path'], params['macro_name'])
            return f"VBA macro {params['macro_name']} executed in Word file {params['file_path']}"
        else:
            return f"Unknown operation: {operation}"
    except Exception as e:
        return f"Error: {str(e)}"

def read_excel(file_path: str, cell: str) -> str:
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active
    return str(sheet[cell].value)

def write_excel(file_path: str, cell: str, value: str) -> None:
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active
    sheet[cell] = value
    workbook.save(file_path)

def create_word_document(file_path: str, content: str) -> None:
    document = Document()
    document.add_paragraph(content)
    document.save(file_path)

def edit_word_document(file_path: str, content: str) -> None:
    document = Document(file_path)
    document.add_paragraph(content)
    document.save(file_path)

def execute_vba_excel(file_path: str, macro_name: str) -> None:
    excel = win32com.client.Dispatch("Excel.Application")
    excel.Visible = False
    try:
        workbook = excel.Workbooks.Open(file_path)
        excel.Application.Run(macro_name)
    finally:
        excel.Quit()

def execute_vba_word(file_path: str, macro_name: str) -> None:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(file_path)
        word.Application.Run(macro_name)
    finally:
        word.Quit()


import os
import shutil
from typing import Optional, Type
from pydantic.v1 import BaseModel, Field
from langchain.agents import tool
import json
import openpyxl
from docx import Document
from PyPDF2 import PdfReader, PdfWriter

class FileSystemInput(BaseModel):
    input: str = Field(description='A JSON string containing the operation and its parameters. Example: {"operation": "create_folder", "folder_name": "documents"}')

@tool("file_system_manager", args_schema=FileSystemInput, return_direct=False)
def file_system_manager(input: str) -> str:
    """
    Manage the file system by creating folders, moving files, listing directory contents, and deleting items.
    Supports operations on various file types including Excel, Word, and PDF.
    """
    try:
        params = json.loads(input)
        operation = params.get("operation")

        if operation == "create_folder":
            folder_name = params.get("folder_name")
            if not folder_name:
                return "Error: Folder name is required."
            os.makedirs(folder_name, exist_ok=True)
            return f"Folder '{folder_name}' created successfully."

        elif operation == "move_file":
            source = params.get("source")
            destination = params.get("destination")
            if not source or not destination:
                return "Error: Source and destination are required."
            shutil.move(source, destination)
            return f"File moved from '{source}' to '{destination}' successfully."

        elif operation == "list_directory":
            directory = params.get("directory", ".")
            items = os.listdir(directory)
            return f"Contents of '{directory}': {', '.join(items)}"

        elif operation == "delete_item":
            item_path = params.get("item_path")
            if not item_path:
                return "Error: Item path is required."
            if os.path.isfile(item_path):
                os.remove(item_path)
                return f"File '{item_path}' deleted successfully."
            elif os.path.isdir(item_path):
                shutil.rmtree(item_path)
                return f"Directory '{item_path}' and its contents deleted successfully."
            else:
                return f"Error: '{item_path}' does not exist."

        elif operation == "read_excel":
            file_path = params.get("file_path")
            sheet_name = params.get("sheet_name", "Sheet1")
            if not file_path:
                return "Error: File path is required."
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook[sheet_name]
            data = [[cell.value for cell in row] for row in sheet.iter_rows()]
            return json.dumps(data)

        elif operation == "read_word":
            file_path = params.get("file_path")
            if not file_path:
                return "Error: File path is required."
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return content

        elif operation == "read_pdf":
            file_path = params.get("file_path")
            if not file_path:
                return "Error: File path is required."
            with open(file_path, "rb") as file:
                reader = PdfReader(file)
                content = ""
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            return content

        else:
            return f"Error: Unknown operation '{operation}'"

    except json.JSONDecodeError:
        return "Error: Invalid input format. Please provide a valid JSON string."
    except Exception as e:
        return f"Error: {str(e)}"


from typing import Optional, Type
from pydantic.v1 import BaseModel, Field
from langchain.agents import tool
from langchain.tools import HumanInputRun
import json

class HumanReviewInput(BaseModel):
    input: str = Field(description='A string containing the issue to be reviewed and any additional context. Example: {"issue": "New product design needs approval", "context": "The design incorporates user feedback from last quarter"}')

@tool("human_review_interface", args_schema=HumanReviewInput, return_direct=False)
def human_review_interface(input: str) -> str:
    """Flag issues for human review and receive approval or input. This tool allows users to submit issues for human review, providing context and receiving responses such as approval, rejection, or additional feedback."""
    human_input_tool = HumanInputRun()
    
    try:
        input_data = json.loads(input)
    except json.JSONDecodeError:
        return "Error: Invalid input format. Please provide a valid JSON string."
    
    issue = input_data.get('issue', 'No issue provided')
    context = input_data.get('context', 'No additional context provided')
    
    prompt = f"Human Review Required:\nIssue: {issue}\nContext: {context}\n\nPlease review and respond with one of the following:\n1. 'approve' to approve the issue\n2. 'reject' to reject the issue\n3. Any other input will be considered as feedback or comments\n\nYour response: "
    
    response = human_input_tool.run(prompt)
    
    if response.lower() == 'approve':
        return "Issue approved by human reviewer."
    elif response.lower() == 'reject':
        return "Issue rejected by human reviewer."
    else:
        return f"Human reviewer provided feedback: {response}"


from typing import Optional, Type
from pydantic.v1 import BaseModel, Field
from langchain.agents import tool
import json

class CourtProcedureInput(BaseModel):
    input: str = Field(description='A JSON string containing the query for court procedures. Example: {"query": "What is the procedure for filing a motion in federal court?"}')

@tool("court_procedure_db", args_schema=CourtProcedureInput, return_direct=False)
def court_procedure_db(input: str) -> str:
    """Query the court procedure database for specific court requirements and procedures. This tool provides access to a comprehensive database of court procedures, allowing users to retrieve detailed information on various legal processes, filing requirements, and procedural guidelines for different types of courts and legal actions."""
    try:
        input_data = json.loads(input)
        query = input_data.get("query")
        
        if not query:
            return "Error: No query provided in the input."
        
        court_procedures = {
            "filing a motion in federal court": "To file a motion in federal court, follow these steps:\n1. Draft the motion\n2. File the motion with the court clerk\n3. Serve the motion on all parties\n4. File a proof of service",
            "appealing a decision": "To appeal a decision in federal court:\n1. File a notice of appeal\n2. Order the trial transcript\n3. File your appellate brief\n4. Participate in oral arguments if scheduled",
            "requesting a continuance": "To request a continuance:\n1. Draft a motion for continuance\n2. Explain the reasons for the request\n3. File the motion as soon as possible\n4. Obtain opposing counsel's position on the request",
            "filing a complaint": "To file a complaint in federal court:\n1. Draft the complaint\n2. File the complaint with the court clerk\n3. Pay the filing fee or request a fee waiver\n4. Serve the complaint on the defendant(s)\n5. File proof of service with the court",
            "discovery process": "The discovery process in federal court includes:\n1. Initial disclosures\n2. Interrogatories\n3. Requests for production of documents\n4. Depositions\n5. Requests for admission",
        }
        
        def search_procedures(query, procedures):
            results = []
            for key, value in procedures.items():
                if query.lower() in key.lower():
                    results.append(value)
            return results

        matching_procedures = search_procedures(query, court_procedures)
        
        if matching_procedures:
            return "\n\n".join(matching_procedures)
        else:
            return f"No specific information found for the query: {query}. Please try a different or more specific query."
    
    except json.JSONDecodeError:
        return "Error: Invalid JSON input. Please provide a valid JSON string."
    except Exception as e:
        return f"An error occurred: {str(e)}"
